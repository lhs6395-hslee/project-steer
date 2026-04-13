"""Generate DOCX template based on DOCX_STYLE_GUIDE.md specifications."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def setup_styles(doc):
    """Configure Normal and Heading styles per style guide."""
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)

    headings = {1: (14, "1A1A2E"), 2: (13, "1A1A2E"), 3: (12, "1A1A2E")}
    for lv, (sz, clr) in headings.items():
        hs = doc.styles[f"Heading {lv}"]
        hs.font.size = Pt(sz)
        hs.font.color.rgb = RGBColor.from_string(clr)
        hs.font.bold = True
        hs.font.name = "맑은 고딕"
        hs.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def add_cover_page(doc):
    """Add cover page with title, subtitle, date, separator, company."""
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(100)  # ~1440 twips
    pf.space_after = Pt(17)    # ~240 twips
    r = p.add_run("보고서 제목")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    r.font.name = "맑은 고딕"

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(11)
    r = p.add_run("부제목을 입력하세요")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    r.font.name = "맑은 고딕"

    # Date with bottom border
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(28)
    r = p.add_run("2026-01-01")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    r.font.name = "맑은 고딕"
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Company
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(150)  # ~2160 twips
    r = p.add_run("회사명")
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = "맑은 고딕"

    # Page break
    doc.add_page_break()


def add_history_page(doc):
    """Add History table on page 2."""
    doc.add_heading("History", level=1)
    headers = ["버전", "일자", "작성자", "내용"]
    rows = [["1.0", "2026-01-01", "작성자명", "초안 작성"]]
    add_styled_table(doc, headers, rows)
    doc.add_page_break()


def add_styled_table(doc, headers, rows):
    """Create a table with navy header and zebra striping."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = "맑은 고딕"
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1B3A5C")
        shd.set(qn("w:val"), "clear")
        c._tc.get_or_add_tcPr().append(shd)

    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            c = t.rows[ri + 1].cells[ci]
            c.text = ""
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            r.font.name = "맑은 고딕"
            if ri % 2 == 1:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F2F2F2")
                shd.set(qn("w:val"), "clear")
                c._tc.get_or_add_tcPr().append(shd)
    return t


def add_footer(doc):
    """Add centered page number footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), " PAGE ")
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "18")
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = "1"
        r.append(t)
        fld.append(r)
        p._p.append(fld)


def add_sample_body(doc):
    """Add sample body structure."""
    doc.add_heading("1. 개요", level=1)
    doc.add_heading("1.1 목적", level=2)
    doc.add_paragraph("본 문서의 목적을 기술한다.")
    doc.add_heading("1.2 범위", level=2)
    doc.add_paragraph("본 문서의 범위를 기술한다.")

    doc.add_page_break()
    doc.add_heading("2. 본문", level=1)
    doc.add_heading("2.1 세부 내용", level=2)
    doc.add_paragraph("세부 내용을 기술한다.")

    # Sample table
    headers = ["항목", "설명", "비고"]
    rows = [
        ["항목 1", "설명 1", "-"],
        ["항목 2", "설명 2", "-"],
    ]
    add_styled_table(doc, headers, rows)
    doc.add_paragraph("위 표와 같이 항목별 세부 내용을 정리한다.")


def main():
    doc = Document()
    setup_styles(doc)
    add_cover_page(doc)
    add_history_page(doc)
    add_sample_body(doc)
    add_footer(doc)
    doc.save("templates/docx_template.docx")
    print("✅ templates/docx_template.docx 생성 완료")


if __name__ == "__main__":
    main()
